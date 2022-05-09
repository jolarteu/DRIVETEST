from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, glob
import os
from freq import *
from django.conf import settings
from django.core.files.storage import default_storage
import ntpath

script_dir = os.path.dirname(__file__)
mapas_dir = os.path.join(script_dir, 'mapas')
graficas_dir = os.path.join(script_dir, 'graficas')

def mapa(latitud, longitud, nombre, type):

    colors=['red','green','blue','black', 'orange','gray','yellow']
    ruh_m = plt.imread('map.png')
    nombre=nombre+".png"
    # BBox = ((-74.1704,  -74.0097,
    #     4.5982, 4.7138))
    BBox = ((-74.2453,  -73.9239,
        4.5412, 4.7725))



    fig, ax = plt.subplots()
    if type==True:
        for i in range(len(latitud)):
            ax.scatter( longitud[i],latitud[i], s=1, zorder=1, alpha= 1, c=colors[i])
    else:
        ax.scatter( longitud,latitud, s=1, zorder=1, alpha= 1, c='b')
    ax.set_title('Drive Test Powered by Sostecnible')
    ax.set_xlim(BBox[0],BBox[1])
    ax.set_ylim(BBox[2],BBox[3])
    legend1 = ax.legend()
    ax.add_artist(legend1)
    ax.imshow(ruh_m, zorder=0, extent = BBox, aspect= 'auto')
    dir = os.path.join(script_dir, 'mapas')
    dir2 = os.path.join(dir, nombre)
    print(dir2)
    plt.savefig(dir2,  dpi=250)

def generar_mapas(df):
    latitud=df['LAT']
    longitud=df['LON']
    lista=[1, 0, 3, 4]
    for i in lista:
        nombre="mapa_"+str(i)
        if i==1:
            mapa(latitud, longitud, nombre, False)
        else:
            lista=df.loc[df['SYSTEM'] ==i]
            latitud=lista['LAT']
            longitud=lista['LON']
            mapa(latitud, longitud, nombre, False)

def generar_histogramas(df):
    lista=[ 'PLMN', 'SYSTEM','BAND']
    for i in lista:
        nombre = os.path.join(graficas_dir, str(i))
        histogram(i, nombre, df)

def generar_mapas_roaming(df):

    a = df['PLMN'].unique()
    latitud=[]
    longitud=[]
    nombre="REDES"
    for i in a:
        # nombre="mapa_"+str(i)
        lista=df.loc[df['PLMN'] ==i]
        latitud.append(lista['LAT'])
        longitud.append(lista['LON'])

    print(len(longitud))
    print(len(latitud))

    mapa(latitud, longitud, nombre, True)

def generar_mapas_roaming_2(df):

    a = df['PLMN'].unique()

    for i in a:
        print (i)
        nombre="mapa_"+str(i)
        lista=df.loc[df['PLMN'] ==i]
        latitud=lista['LAT']
        longitud=lista['LON']
        mapa(latitud, longitud, nombre, False)

def generar_mapas_ca(df):
    lista=df.loc[df['CA'] ==2]
    latitud=lista['LAT']
    longitud=lista['LON']
    nombre="mapa_CA"
    # latitud=[]
    # longitud=[]
    # for i in a:
    #     print (i)
    #     lista=df.loc[df['CA'] ==i]
    #     latitud.append(lista['LAT'])
    #     longitud.append(lista['LON'])

    mapa(latitud, longitud, nombre, False)

def guardar():

    document = Document()
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture("sos.png")
    lista=['PLMN', 'SYSTEM','BAND']
    puntos=[".0",""]

    for i in lista:
        nombre = os.path.join(graficas_dir, str(i))

        document.add_paragraph(
        'histograma '+str(i), style='List Bullet'
        )

        document.add_picture(nombre+".png", width=Inches(5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


    for i, j in zip([1, 4,3,0], ['mapa recorrido', 'mapa 4G', 'mapa 3G','mapa sin cobertura']):

        mapa=os.path.join(mapas_dir, 'mapa_'+str(i)+'.png')

        if not os.path.isfile(mapa):
            continue

        document.add_paragraph(
        j, style='List Bullet'
        )
        document.add_picture(mapa, width=Inches(5))
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    mapa_redes= os.path.join(mapas_dir, 'REDES.png')

    document.add_paragraph(
    'Mapa red redes', style='List Bullet'
    )
    document.add_picture(mapa_redes, width=Inches(5))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #
    for punto in puntos:

        mapa_360= os.path.join(mapas_dir, 'mapa_732360'+punto+'.png')
        if not os.path.isfile(mapa_360):
            continue

        for i, j in zip([732360,732130,732123,732103,732101]  , ['mapa red 360', 'mapa red 130', 'mapa red 123','mapa red 103', 'mapa red 101']):

            mapa= os.path.join(mapas_dir, 'mapa_'+str(i)+punto+'.png')
            if not os.path.isfile(mapa):
                continue
            document.add_paragraph(
            j, style='List Bullet'
            )

            document.add_picture(mapa, width=Inches(5))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save('demo.docx')

def borrar():
    dir = os.path.join(script_dir, 'mapas')
    # dir = "/Users/juans/Documents/informe_drivetest/mapas"
    for file in os.scandir(dir):
         os.remove(file.path)

def main_map(df):
    df.replace(os.sep,ntpath.sep)
    df=df[1:]
    df=pd.read_csv(df,  sep=';')
    borrar()
    gps_(df,graficas_dir)
    generar_histogramas(df)
    # #generar_mapas_ca(df)
    generar_mapas(df)
    generar_mapas_roaming(df)
    generar_mapas_roaming_2(df)
    # #generar_mapas_ca(df)
    guardar()

    #borrar()

#
# if __name__ == '__main__':
#     #revisar()
#     main()
